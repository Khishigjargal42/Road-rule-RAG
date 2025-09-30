#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
DOCX/JSON -> Chroma indexer with pluggable embeddings (Ollama / OpenAI / HuggingFace),
batching, optional full reindex, and simple progress logging.

Examples
--------
# 1) Local Ollama (fast local embed model)
python rag_chroma_langchain_upgraded.py --json road_rules_structured.json \
  --backend ollama --embed-model mxbai-embed-large --persist ./chroma_db --collection road_rules \
  --batch-size 64 --progress

# 2) OpenAI (requires OPENAI_API_KEY in env)
python rag_chroma_langchain_upgraded.py --json road_rules_structured.json \
  --backend openai --embed-model text-embedding-3-small --clear --progress

# 3) HuggingFace (offline-friendly)
python rag_chroma_langchain_upgraded.py --json road_rules_structured.json \
  --backend hf --embed-model sentence-transformers/all-MiniLM-L6-v2 --batch-size 64 --progress
"""
from __future__ import annotations

import os
import re
import json
import argparse
from typing import List, Tuple, Dict, Any
import hashlib
from pathlib import Path
import shutil

# --- LangChain & Chroma ---
from langchain_chroma import Chroma
from langchain_core.documents import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter

# Embedding backends (import lazily inside factory where possible)
from langchain_ollama import OllamaEmbeddings  # for --backend ollama
try:
    from langchain_openai import OpenAIEmbeddings  # for --backend openai
except Exception:
    OpenAIEmbeddings = None
try:
    from langchain_community.embeddings import HuggingFaceEmbeddings  # for --backend hf
except Exception:
    HuggingFaceEmbeddings = None

# --- Optional DOCX parsing ---
try:
    from docx import Document as DocxDocument  # pip install python-docx
except Exception:
    DocxDocument = None

# -----------------------
# Helpers
# -----------------------
CHAPTER_WORDS = r"(Нэг|Хоёр|Гурав|Дөрөв|Тав|Зургаа|Долоо|Найм|Ес|Арав)"
chapter_re = re.compile(rf"^{CHAPTER_WORDS}\.\s+(.+)$")
article_re = re.compile(r"^(\d+(?:\.\d+)+)\.\s+(.*)$")
penalty_re = re.compile(r"(торгуу|шийтгэл)|₮|\b(мянга|сая)\s*төгрөг", re.IGNORECASE)


def sha1_16(s: str) -> str:
    return hashlib.sha1(s.encode("utf-8", errors="ignore")).hexdigest()[:16]


def make_block_id(chapter: str | None, article: str | None, text: str) -> str:
    return sha1_16(f"{chapter or ''}||{article or ''}||{text.strip()}")


# -----------------------
# Load & normalize JSON (supports 2 schemas)
# -----------------------

def load_blocks(json_path: str) -> Tuple[str, List[Dict[str, Any]]]:
    """
    Supports either:
    A) { "meta": {...}, "content": [ {type: heading|paragraph|list|table, ...}, ... ] }
    B) [ { "chapter": "...", "article": "1.2.1", "text": "...", "type": "definition|rule|penalty" }, ... ]
    Returns (filename, blocks as a unified list[dict])
    """
    with open(json_path, "r", encoding="utf-8") as f:
        data = json.load(f)

    filename = os.path.basename(json_path)

    # Schema A (meta+content)
    if isinstance(data, dict) and "content" in data:
        blocks = data.get("content", [])
        # Ensure minimal shape
        norm = []
        for b in blocks:
            t = (b.get("type") or "").lower()
            if t in {"heading", "paragraph", "list", "table"}:
                norm.append(b)
        return filename, norm

    # Schema B (list of records)
    if isinstance(data, list):
        # Convert records -> heading/paragraph blocks (group by chapter/article)
        blocks = []
        seen_chapters = set()
        for r in data:
            chap = r.get("chapter") or ""
            if chap and chap not in seen_chapters:
                seen_chapters.add(chap)
                blocks.append({"type": "heading", "text": chap})
            # article as sub-heading to help chunking
            art = r.get("article")
            if art:
                blocks.append({"type": "heading", "text": f"Заалт {art}"})
            txt = (r.get("text") or "").strip()
            if txt:
                blocks.append({"type": "paragraph", "text": txt})
        return filename, blocks

    raise ValueError("Unrecognized JSON schema. Provide either {meta, content[]} or list-of-records.")


# -----------------------
# DOCX -> records (rules-style)
# -----------------------

def parse_docx_rules(docx_path: str) -> List[Dict[str, Any]]:
    if DocxDocument is None:
        raise RuntimeError("python-docx is not installed. pip install python-docx")
    doc = DocxDocument(docx_path)
    chapter = None
    records: List[Dict[str, Any]] = []
    buffer_article_id, buffer_text_parts = None, []

    def flush():
        nonlocal buffer_article_id, buffer_text_parts, chapter, records
        if buffer_article_id is None:
            return
        text = " ".join([t.strip() for t in buffer_text_parts]).strip()
        rec_type = "definition" if ("“" in text and "”" in text) else "rule"
        if penalty_re.search(text) and rec_type != "definition":
            rec_type = "penalty"
        rec = {
            "chapter": chapter,
            "article": buffer_article_id,
            "type": rec_type,
            "text": text,
        }
        m = re.search(r"“([^”]+)”", text)
        if m:
            rec["term"] = m.group(1)
        records.append(rec)
        buffer_article_id, buffer_text_parts = None, []

    for p in doc.paragraphs:
        ln = p.text.strip()
        if not ln:
            continue
        mch = chapter_re.match(ln)
        if mch:
            flush()
            chapter = f"{mch.group(0)}"  # keep entire "Нэг. ..."
            continue
        ma = article_re.match(ln)
        if ma:
            flush()
            buffer_article_id = ma.group(1)
            buffer_text_parts = [ma.group(2)]
        else:
            if buffer_article_id is not None:
                buffer_text_parts.append(ln)
    flush()
    return records


# -----------------------
# Build LangChain Documents with chunking
# -----------------------

def build_documents_from_blocks(filename: str, blocks: List[Dict[str, Any]], chunk_size=1000, chunk_overlap=150) -> Tuple[List[Document], List[str]]:
    splitter = RecursiveCharacterTextSplitter(chunk_size=chunk_size, chunk_overlap=chunk_overlap)
    docs: List[Document] = []
    ids: List[str] = []

    section_title = "Introduction"
    buffer_lines: List[str] = []

    def flush_section():
        nonlocal buffer_lines, section_title, docs, ids
        if not buffer_lines:
            return
        text = "\n".join(buffer_lines).strip()
        if not text:
            buffer_lines.clear()
            return
        chunks = splitter.split_text(text)
        start_index = len(docs)
        for j, ch in enumerate(chunks):
            meta = {
                "source": filename,
                "section": section_title,
                "chunk": start_index + j,
            }
            docs.append(Document(page_content=ch, metadata=meta))
            ids.append(f"{filename}::section={section_title}::chunk={start_index + j}")
        buffer_lines.clear()

    for b in blocks:
        btype = (b.get("type") or "").lower()
        if btype == "heading":
            flush_section()
            section_title = (b.get("text") or "Untitled").strip() or "Untitled"
        elif btype == "paragraph":
            t = (b.get("text") or "").strip()
            if t:
                buffer_lines.append(t)
        elif btype == "list":
            for item in b.get("items", []):
                t = (item.get("text") or "").strip()
                if t:
                    buffer_lines.append(t)
        elif btype == "table":
            for row in b.get("rows", []):
                line = " | ".join([c for c in row if c])
                if line.strip():
                    buffer_lines.append(line)
        else:
            # ignore others
            pass
    flush_section()
    return docs, ids


def build_documents_from_records(filename: str, records: List[Dict[str, Any]], chunk_size=800, chunk_overlap=100) -> Tuple[List[Document], List[str]]:
    """
    Alternate path: directly chunk each record's text, keeping chapter/article in metadata.
    """
    splitter = RecursiveCharacterTextSplitter(chunk_size=chunk_size, chunk_overlap=chunk_overlap)
    docs, ids = [], []
    for r in records:
        text = (r.get("text") or "").strip()
        if not text:
            continue
        chunks = splitter.split_text(text)
        for j, ch in enumerate(chunks):
            meta = {
                "source": filename,
                "chapter": r.get("chapter"),
                "article": r.get("article"),
                "type": r.get("type"),
                "term": r.get("term"),
                "chunk": j,
            }
            docs.append(Document(page_content=ch, metadata=meta))
            ids.append(make_block_id(r.get("chapter"), r.get("article"), f"{ch}::{j}"))
    return docs, ids


# -----------------------
# Embedding backends
# -----------------------

def build_embeddings(backend: str, model_name: str):
    backend = (backend or "ollama").lower()
    if backend == "ollama":
        return OllamaEmbeddings(model=model_name)
    if backend == "openai":
        if OpenAIEmbeddings is None:
            raise RuntimeError("langchain_openai not installed. pip install langchain-openai")
        if not os.getenv("OPENAI_API_KEY"):
            raise RuntimeError("Set OPENAI_API_KEY in your environment for --backend openai.")
        return OpenAIEmbeddings(model=model_name)
    if backend == "hf":
        if HuggingFaceEmbeddings is None:
            raise RuntimeError("langchain-community not installed. pip install langchain-community sentence-transformers")
        # Typical small, fast default if user didn't pass a specific model
        model_name = model_name or "sentence-transformers/all-MiniLM-L6-v2"
        return HuggingFaceEmbeddings(model_name=model_name)
    raise ValueError(f"Unknown backend: {backend}")


def get_vector_store(collection: str, persist_dir: str, embeddings) -> Chroma:
    return Chroma(
        collection_name=collection,
        persist_directory=persist_dir,
        embedding_function=embeddings,
    )


def batched(seq, batch_size: int):
    n = len(seq)
    for i in range(0, n, batch_size):
        yield i, seq[i : i + batch_size]


def index_into_chroma(
    docs: List[Document],
    ids: List[str],
    collection: str,
    persist: str,
    backend: str,
    embed_model: str,
    clear: bool,
    batch_size: int,
    progress: bool,
):
    # Optionally clear old DB to avoid dim mismatch
    if clear and Path(persist).exists():
        shutil.rmtree(persist, ignore_errors=True)

    embeddings = build_embeddings(backend, embed_model)
    vs = get_vector_store(collection, persist, embeddings)

    print(f"[Index] Upserting {len(ids)} chunks into '{collection}' at '{persist}' using backend={backend} model={embed_model} ...")

    # Try fast path: precompute embeddings in batches then add_embeddings (skips re-embed inside Chroma)
    texts = [d.page_content for d in docs]
    metadatas = [d.metadata for d in docs]

    vectors: List[List[float]] = []
    if batch_size < 1:
        batch_size = 16

    for i, batch_texts in batched(texts, batch_size):
        vecs = embeddings.embed_documents(batch_texts)
        vectors.extend(vecs)
        if progress:
            print(f"  [embed] {min(i+len(batch_texts), len(texts))}/{len(texts)}")

    # Some versions expose add_embeddings; if not, fall back to add_documents
    try:
        vs.add_embeddings(texts=texts, embeddings=vectors, metadatas=metadatas, ids=ids)
    except Exception:
        # Fall back (this will re-embed inside VS, but our precompute didn't go to waste—kept as safety)
        if progress:
            print("  [warn] add_embeddings not available; falling back to add_documents (may be slower)")
        vs.add_documents(documents=docs, ids=ids)

    print("[Index] Done.")
    return vs


# -----------------------
# Retrieval quick test
# -----------------------

def quick_query(vs: Chroma, q: str, k: int = 5):
    retriever = vs.as_retriever(search_kwargs={"k": k})
    print("\n[Query]", q)
    results = retriever.get_relevant_documents(q)
    for i, r in enumerate(results, 1):
        print("----", i)
        print(r.metadata)
        print(r.page_content[:300].replace("\n", " "), "...")


# -----------------------
# Main CLI
# -----------------------

def main():
    ap = argparse.ArgumentParser(description="DOCX/JSON -> Chroma with pluggable embeddings (LangChain pipeline).")
    g = ap.add_mutually_exclusive_group(required=True)
    g.add_argument("--docx", help="Path to a rules-style .docx")
    g.add_argument("--json", help="Path to JSON (either meta+content or list-of-records)")

    ap.add_argument("--persist", default="./chroma_db", help="Chroma persist directory")
    ap.add_argument("--collection", default="road_rules", help="Chroma collection name")

    ap.add_argument("--backend", default="ollama", choices=["ollama", "openai", "hf"], help="Embedding backend")
    ap.add_argument("--embed-model", default="mxbai-embed-large", help="Embedding model name for the chosen backend")

    ap.add_argument("--clear", action="store_true", help="Delete persist dir before indexing (avoid dim mismatch, clean start)")
    ap.add_argument("--batch-size", type=int, default=64, help="Batch size for embedding (precompute phase)")
    ap.add_argument("--progress", action="store_true", help="Log simple progress while embedding")

    ap.add_argument("--reindex", action="store_true", help="(Deprecated) alias for --clear")
    ap.add_argument("--query", default=None, help="Run a quick retrieval query after indexing")
    ap.add_argument("--record-chunk", action="store_true", help="Chunk per record instead of section-based")

    args = ap.parse_args()
    if args.reindex:
        args.clear = True

    # Gentle warning for OneDrive paths (can slow down DuckDB/Parquet writes)
    if "OneDrive" in str(Path(args.persist).resolve()):
        print("[Warn] Persist dir is under OneDrive; consider moving to a non-synced path for speed.")

    if args.docx and DocxDocument is None:
        raise RuntimeError("python-docx not installed. pip install python-docx")

    # Prepare docs + ids
    if args.docx:
        print(f"[Parse] DOCX -> records: {args.docx}")
        records = parse_docx_rules(args.docx)
        filename = os.path.basename(args.docx)
        if args.record_chunk:
            docs, ids = build_documents_from_records(filename, records)
        else:
            # Convert via blocks for better section grouping
            # Turn records into blocks
            blocks: List[Dict[str, Any]] = []
            seen_chapters = set()
            for r in records:
                chap = r.get("chapter")
                if chap and chap not in seen_chapters:
                    seen_chapters.add(chap)
                    blocks.append({"type": "heading", "text": chap})
                if r.get("article"):
                    blocks.append({"type": "heading", "text": f"Заалт {r['article']}`"})
                if r.get("text"):
                    blocks.append({"type": "paragraph", "text": r["text"]})
            docs, ids = build_documents_from_blocks(filename, blocks)
    else:
        filename, blocks_or_records = load_blocks(args.json)
        # Heuristic: if items have 'type' heading/paragraph -> treat as blocks
        if blocks_or_records and isinstance(blocks_or_records[0], dict) and blocks_or_records[0].get("type") in {"heading", "paragraph", "list", "table"}:
            docs, ids = build_documents_from_blocks(filename, blocks_or_records)
        else:
            # Already normalized to blocks in loader, but keep both paths
            docs, ids = build_documents_from_blocks(filename, blocks_or_records)

    print(f"[Docs] Prepared {len(docs)} chunks.")

    vs = index_into_chroma(
        docs=docs,
        ids=ids,
        collection=args.collection,
        persist=args.persist,
        backend=args.backend,
        embed_model=args.embed_model,
        clear=args.clear,
        batch_size=args.batch_size,
        progress=args.progress,
    )

    if args.query:
        quick_query(vs, args.query)


if __name__ == "__main__":
    main()
